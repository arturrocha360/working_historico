from docx import Document
from docx.shared import Pt
from docx.enum.text import  WD_PARAGRAPH_ALIGNMENT
from tkinter import messagebox
import sqlite3
from Caminhos_documentos import caminho_template_lista_ip,caminho_tabela_plano_integração



def ip_consulta(estacao, Descricao_Sistema):
    # Conexão com o banco de dados SQLite3
    conexao = sqlite3.connect('dados.db')

    # Cursor para executar comandos SQL
    cursor = conexao.cursor()
    # Consulta SQL com filtro pelo nome
    consulta = f"""
            SELECT *
            FROM tabela_ip
            WHERE Estação = '{estacao}' AND
                Descrição_Sistema = '{Descricao_Sistema}'   
                    
            """
    # Executar a consulta
    cursor.execute(consulta)

    # Recuperar os resultados da consulta
    resultados = cursor.fetchall()

    # Fechar a conexão
    conexao.close()
    return resultados

def criar_wordip(lista_recebida):
    # Carregue o documento existente
    doc = Document(caminho_template_lista_ip)
    # Acesse a tabela existente (no caso do documento lista de IPs a tabela 5 (4 em python) é a tabela que deve adicionar conteúdo)
    tabela = doc.tables[4]

    # Adicione conteúdo a células específicas da tabela
    
    for i in lista_recebida:
        nova_linha = tabela.add_row().cells
        nova_linha[0].text = i[0]
        nova_linha[1].text = i[1]
        nova_linha[2].text = i[2]
        nova_linha[3].text = i[3]
        nova_linha[4].text = i[4]
        nova_linha[5].text = i[5]
        nova_linha[6].text = i[6]

    fonte = 'Arial'  # Substitua 'Arial' pela fonte desejada
    tamanho = Pt(7)  # Substitua 12 pelo tamanho de fonte desejado

    tabela.style = 'Table Grid'

    for row in tabela.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = fonte
                    run.font.size = tamanho

    # Salve o documento atualizado
    doc.save(caminho_template_lista_ip)
    
    messagebox.showinfo("Aviso", "Documento Lista de IP's Gerado!")

def gerar_listaip(tag_,descricao_,ip_,mascara_,gateway_,rede_,id_vlan_):
    
    matriz=[]
    lista=[tag_,descricao_,ip_,mascara_,gateway_,rede_,id_vlan_]
    for i in lista:
        matriz.append(i)
    matriz_invertida=list(zip(*matriz))
    criar_wordip(matriz_invertida)

def item8_plan_int(lista_recebida):
    # Carregue o documento existente
    doc = Document(caminho_tabela_plano_integração)
    # Acesse a tabela existente (no caso do documento lista de IPs a tabela 5 (4 em python) é a tabela que deve adicionar conteúdo)
    tabela = doc.tables[-1]
    invertido = list(zip(*lista_recebida))
    # Adicione conteúdo a células específicas da tabela
    
    for i in invertido:
        nova_linha = tabela.add_row().cells
        nova_linha[0].text = i[0]
        nova_linha[1].text = i[1]
        nova_linha[2].text = i[2]
        nova_linha[3].text = i[3]
        
    fonte = 'Arial'  # Substitua 'Arial' pela fonte desejada
    tamanho = Pt(7)  # Substitua 12 pelo tamanho de fonte desejado

    tabela.style = 'Table Grid'

    for row in tabela.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = fonte
                    run.font.size = tamanho

    # Salve o documento atualizado
    doc.save(caminho_tabela_plano_integração)
